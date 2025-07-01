import os
import tempfile
import uuid
from datetime import datetime, timedelta
from flask import Flask, request, jsonify, send_file, session, redirect, url_for
from flask_cors import CORS
from werkzeug.utils import secure_filename
import azure.cognitiveservices.speech as speechsdk
from google_auth_oauthlib.flow import Flow
from google.auth.transport.requests import Request
from google.oauth2 import id_token
import google.auth.transport.requests
import requests
import io
from pydub import AudioSegment
from docx import Document
import json
from functools import wraps
import stripe
import sqlite3
from contextlib import contextmanager

# Try to import PostgreSQL support, fall back to SQLite if not available
try:
    import psycopg2
    from psycopg2.extras import RealDictCursor
    POSTGRES_AVAILABLE = True
except ImportError:
    POSTGRES_AVAILABLE = False
    print("psycopg2 not available, using SQLite for database operations")

# Disable OAuth2 HTTPS requirement for local development
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'

app = Flask(__name__)
CORS(app, supports_credentials=True, origins=['http://localhost:5000'])
app.secret_key = os.environ.get('FLASK_SECRET_KEY', 'dev-secret-key-change-in-production-seriously')

# Session configuration
app.config['SESSION_COOKIE_SECURE'] = False
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'

# Application Configuration
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'wav', 'mp3', 'aac', 'm4a', 'flac', 'ogg'}

# Environment variables
GOOGLE_CLIENT_ID = os.environ.get('GOOGLE_CLIENT_ID')
GOOGLE_CLIENT_SECRET = os.environ.get('GOOGLE_CLIENT_SECRET')
AZURE_SPEECH_KEY = os.environ.get('AZURE_SPEECH_KEY')
AZURE_SPEECH_REGION = os.environ.get('AZURE_SPEECH_REGION')
AZURE_SPEECH_ENDPOINT = os.environ.get('AZURE_SPEECH_ENDPOINT')

# Stripe Configuration
stripe.api_key = os.environ.get('STRIPE_SECRET_KEY')
STRIPE_PUBLISHABLE_KEY = os.environ.get('STRIPE_PUBLISHABLE_KEY')
STRIPE_WEBHOOK_SECRET = os.environ.get('STRIPE_WEBHOOK_SECRET')
SUBSCRIPTION_PRICE_ID = os.environ.get('STRIPE_PRICE_ID')  # Your £5/month price ID

# Database setup
DATABASE_URL = os.environ.get('DATABASE_URL', 'sqlite:///app.db')

def init_db():
    """Initialize the database with required tables."""
    if DATABASE_URL.startswith('postgresql://') and POSTGRES_AVAILABLE:
        # PostgreSQL initialization
        try:
            conn = psycopg2.connect(DATABASE_URL)
            with conn.cursor() as cur:
                cur.execute('''
                    CREATE TABLE IF NOT EXISTS users (
                        id SERIAL PRIMARY KEY,
                        google_id TEXT UNIQUE NOT NULL,
                        email TEXT NOT NULL,
                        name TEXT NOT NULL,
                        picture TEXT,
                        stripe_customer_id TEXT,
                        subscription_status TEXT DEFAULT 'inactive',
                        subscription_id TEXT,
                        subscription_end_date TIMESTAMP,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                    )
                ''')
                
                cur.execute('''
                    CREATE TABLE IF NOT EXISTS transcriptions (
                        id SERIAL PRIMARY KEY,
                        user_id INTEGER NOT NULL,
                        filename TEXT NOT NULL,
                        transcript TEXT NOT NULL,
                        word_count INTEGER,
                        confidence REAL,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        FOREIGN KEY (user_id) REFERENCES users (id)
                    )
                ''')
            conn.commit()
            conn.close()
            print("PostgreSQL database initialized successfully")
        except Exception as e:
            print(f"PostgreSQL initialization failed: {e}")
    else:
        # SQLite fallback
        database_file = DATABASE_URL.replace('sqlite:///', '') if DATABASE_URL.startswith('sqlite://') else 'app.db'
        with sqlite3.connect(database_file) as conn:
            conn.execute('''
                CREATE TABLE IF NOT EXISTS users (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    google_id TEXT UNIQUE NOT NULL,
                    email TEXT NOT NULL,
                    name TEXT NOT NULL,
                    picture TEXT,
                    stripe_customer_id TEXT,
                    subscription_status TEXT DEFAULT 'inactive',
                    subscription_id TEXT,
                    subscription_end_date DATETIME,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    updated_at DATETIME DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            
            conn.execute('''
                CREATE TABLE IF NOT EXISTS transcriptions (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_id INTEGER NOT NULL,
                    filename TEXT NOT NULL,
                    transcript TEXT NOT NULL,
                    word_count INTEGER,
                    confidence REAL,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (user_id) REFERENCES users (id)
                )
            ''')
        print("SQLite database initialized successfully")

@contextmanager
def get_db_connection():
    """Context manager for database connections."""
    if DATABASE_URL.startswith('postgresql://') and POSTGRES_AVAILABLE:
        # PostgreSQL connection
        conn = psycopg2.connect(DATABASE_URL, cursor_factory=RealDictCursor)
        try:
            yield conn
        finally:
            conn.close()
    else:
        # SQLite fallback
        database_file = DATABASE_URL.replace('sqlite:///', '') if DATABASE_URL.startswith('sqlite://') else 'app.db'
        conn = sqlite3.connect(database_file)
        conn.row_factory = sqlite3.Row
        try:
            yield conn
        finally:
            conn.close()

def execute_query(conn, query, params=None, fetch=False):
    """Execute query with proper syntax for PostgreSQL or SQLite."""
    if DATABASE_URL.startswith('postgresql://') and POSTGRES_AVAILABLE:
        # PostgreSQL uses %s placeholders
        pg_query = query.replace('?', '%s')
        with conn.cursor() as cur:
            cur.execute(pg_query, params or ())
            if fetch == 'one':
                return cur.fetchone()
            elif fetch == 'all':
                return cur.fetchall()
            conn.commit()
    else:
        # SQLite uses ? placeholders
        if fetch == 'one':
            return conn.execute(query, params or ()).fetchone()
        elif fetch == 'all':
            return conn.execute(query, params or ()).fetchall()
        else:
            conn.execute(query, params or ())
            conn.commit()

def get_or_create_user(google_user_info):
    """Get or create user in database."""
    with get_db_connection() as conn:
        user = execute_query(
            conn,
            'SELECT * FROM users WHERE google_id = ?',
            (google_user_info['id'],),
            fetch='one'
        )
        
        if user:
            # Update user info
            execute_query(conn, '''
                UPDATE users 
                SET email = ?, name = ?, picture = ?, updated_at = CURRENT_TIMESTAMP
                WHERE google_id = ?
            ''', (google_user_info['email'], google_user_info['name'], 
                  google_user_info.get('picture', ''), google_user_info['id']))
            return dict(user)
        else:
            # Create new user
            if DATABASE_URL.startswith('postgresql://') and POSTGRES_AVAILABLE:
                with conn.cursor() as cur:
                    cur.execute('''
                        INSERT INTO users (google_id, email, name, picture)
                        VALUES (%s, %s, %s, %s)
                        RETURNING id
                    ''', (google_user_info['id'], google_user_info['email'], 
                          google_user_info['name'], google_user_info.get('picture', '')))
                    user_id = cur.fetchone()['id']
                    conn.commit()
            else:
                cursor = conn.execute('''
                    INSERT INTO users (google_id, email, name, picture)
                    VALUES (?, ?, ?, ?)
                ''', (google_user_info['id'], google_user_info['email'], 
                      google_user_info['name'], google_user_info.get('picture', '')))
                conn.commit()
                user_id = cursor.lastrowid
            
            return {
                'id': user_id,
                'google_id': google_user_info['id'],
                'email': google_user_info['email'],
                'name': google_user_info['name'],
                'picture': google_user_info.get('picture', ''),
                'stripe_customer_id': None,
                'subscription_status': 'inactive'
            }

def update_user_subscription(user_id, stripe_customer_id, subscription_status, subscription_id=None, end_date=None):
    """Update user subscription information and refresh session."""
    with get_db_connection() as conn:
        # Ensure end_date is a datetime object if not None
        if isinstance(end_date, str):
            try:
                end_date = datetime.fromisoformat(end_date.replace('Z', '+00:00'))
            except ValueError:
                print(f"Warning: Could not parse end_date string '{end_date}'. Storing as None.")
                end_date = None
        
        execute_query(conn, '''
            UPDATE users 
            SET stripe_customer_id = ?, subscription_status = ?, 
                subscription_id = ?, subscription_end_date = ?, updated_at = CURRENT_TIMESTAMP
            WHERE id = ?
        ''', (stripe_customer_id, subscription_status, subscription_id, end_date, user_id))

        print(f"update_user_subscription: Updated user {user_id} - customer_id: {stripe_customer_id}, status: {subscription_status}, sub_id: {subscription_id}, end_date: {end_date}")

        # Update session if this user is currently logged in
        try:
            if 'user_db' in session and session['user_db']['id'] == user_id:
                # Re-fetch the user to ensure session is fresh
                updated_user = execute_query(conn, 'SELECT * FROM users WHERE id = ?', (user_id,), fetch='one')
                if updated_user:
                    session['user_db'] = dict(updated_user)
                    print(f"update_user_subscription: Session refreshed for user {user_id}")
        except Exception as e:
            print(f"update_user_subscription: Could not update session: {e}")
            # This is expected in webhook context where there's no session

def is_subscription_active(user_id):
    """Check if user has an active subscription."""
    with get_db_connection() as conn:
        user = execute_query(conn,
            'SELECT subscription_status, subscription_end_date FROM users WHERE id = ?',
            (user_id,),
            fetch='one'
        )
        
        if not user:
            return False
            
        if user['subscription_status'] == 'active':
            # Check if subscription hasn't expired
            if user['subscription_end_date']:
                if isinstance(user['subscription_end_date'], str):
                    end_date = datetime.fromisoformat(user['subscription_end_date'].replace('Z', '+00:00'))
                else:
                    end_date = user['subscription_end_date']
                return datetime.now() < end_date
            return True # If no end_date, assume perpetual active (unlikely for subscriptions but defensively coded)
            
        return False

# Stripe webhook handlers
def handle_subscription_created(subscription):
    """Handle subscription created/updated events."""
    try:
        customer_id = subscription['customer']
        subscription_id = subscription['id']
        status = subscription['status']
        
        print(f"handle_subscription_created: Processing subscription {subscription_id} for customer {customer_id} with status {status}")
        
        with get_db_connection() as conn:
            user = execute_query(conn,
                'SELECT id FROM users WHERE stripe_customer_id = ?',
                (customer_id,),
                fetch='one'
            )
            
            if user:
                user_id = user['id']
                # Convert Unix timestamp to datetime object
                end_date = datetime.fromtimestamp(subscription['current_period_end'])
                
                print(f"handle_subscription_created: Found user {user_id}. Updating subscription to status='{status}', id='{subscription_id}', end_date='{end_date}'")
                
                # Update user subscription
                update_user_subscription(user_id, customer_id, status, subscription_id, end_date)
                
                print(f"handle_subscription_created: Successfully updated subscription for user {user_id}")
            else:
                print(f"handle_subscription_created: No user found in DB for Stripe customer ID: {customer_id}")
                # Try to find user by email from Stripe customer
                try:
                    stripe_customer = stripe.Customer.retrieve(customer_id)
                    if stripe_customer.email:
                        user_by_email = execute_query(conn,
                            'SELECT id FROM users WHERE email = ?',
                            (stripe_customer.email,),
                            fetch='one'
                        )
                        if user_by_email:
                            user_id = user_by_email['id']
                            end_date = datetime.fromtimestamp(subscription['current_period_end'])
                            update_user_subscription(user_id, customer_id, status, subscription_id, end_date)
                            print(f"handle_subscription_created: Linked subscription to existing user {user_id} by email")
                        else:
                            print(f"handle_subscription_created: No user found with email {stripe_customer.email}")
                    else:
                        print(f"handle_subscription_created: Stripe customer {customer_id} has no email")
                except Exception as e:
                    print(f"handle_subscription_created: Error retrieving Stripe customer {customer_id}: {e}")
                
    except Exception as e:
        print(f"Error in handle_subscription_created for subscription {subscription.get('id', 'N/A')}: {e}")
        import traceback
        traceback.print_exc()
        raise

def handle_subscription_updated(subscription):
    """Handle subscription updated events."""
    print(f"handle_subscription_updated: Processing subscription update for {subscription.get('id', 'N/A')}")
    handle_subscription_created(subscription)  # Same logic as creation

def handle_subscription_deleted(subscription):
    """Handle subscription deleted/cancelled events."""
    try:
        customer_id = subscription['customer']
        subscription_id = subscription['id']
        
        print(f"handle_subscription_deleted: Processing subscription deletion for customer {customer_id}, subscription {subscription_id}")
        
        with get_db_connection() as conn:
            user = execute_query(conn, 
                'SELECT id, subscription_id FROM users WHERE stripe_customer_id = ?', 
                (customer_id,), 
                fetch='one'
            )
            if user:
                # Only update if this is the current subscription
                if user.get('subscription_id') == subscription_id or not user.get('subscription_id'):
                    update_user_subscription(user['id'], customer_id, 'canceled', None, None)
                    print(f"handle_subscription_deleted: Canceled subscription for customer {customer_id}")
                else:
                    print(f"handle_subscription_deleted: Subscription {subscription_id} does not match current user subscription {user.get('subscription_id')}")
            else:
                print(f"handle_subscription_deleted: No user found for customer {customer_id}")
            
    except Exception as e:
        print(f"Error in handle_subscription_deleted: {e}")
        import traceback
        traceback.print_exc()
        raise

def handle_payment_succeeded(invoice):
    """Handle successful payment events."""
    try:
        if invoice.get('subscription'):
            # Payment for a subscription
            subscription_id = invoice['subscription']
            print(f"handle_payment_succeeded: Payment succeeded for subscription {subscription_id}")
            
            # Fetch the subscription to get updated info
            subscription = stripe.Subscription.retrieve(subscription_id)
            handle_subscription_created(subscription)
        else:
            print(f"handle_payment_succeeded: Payment succeeded for non-subscription invoice {invoice.get('id', 'N/A')}")
            
    except Exception as e:
        print(f"Error in handle_payment_succeeded: {e}")
        import traceback
        traceback.print_exc()
        raise

def handle_payment_failed(invoice):
    """Handle failed payment events."""
    try:
        customer_id = invoice['customer']
        
        print(f"handle_payment_failed: Payment failed for customer {customer_id}")
        
        with get_db_connection() as conn:
            user = execute_query(conn, 
                'SELECT id, subscription_status FROM users WHERE stripe_customer_id = ?', 
                (customer_id,), 
                fetch='one'
            )
            if user:
                # Update status to past_due but keep other subscription info
                current_status = user.get('subscription_status', 'inactive')
                if current_status == 'active':
                    execute_query(conn, '''
                        UPDATE users 
                        SET subscription_status = ?, updated_at = CURRENT_TIMESTAMP
                        WHERE id = ?
                    ''', ('past_due', user['id']))
                    
                    # Update session if this user is currently logged in
                    try:
                        if 'user_db' in session and session['user_db']['id'] == user['id']:
                            session['user_db']['subscription_status'] = 'past_due'
                    except:
                        pass  # Session might not be available in webhook context
                    
                    print(f"handle_payment_failed: Updated user {user['id']} status to 'past_due'")
                else:
                    print(f"handle_payment_failed: User {user['id']} already has status '{current_status}', not updating")
            else:
                print(f"handle_payment_failed: No user found for customer {customer_id}")
            
    except Exception as e:
        print(f"Error in handle_payment_failed: {e}")
        import traceback
        traceback.print_exc()
        raise

# OAuth and Google setup
client_config = None
if GOOGLE_CLIENT_ID and GOOGLE_CLIENT_SECRET:
    client_config = {
        "web": {
            "client_id": GOOGLE_CLIENT_ID,
            "client_secret": GOOGLE_CLIENT_SECRET,
            "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            "token_uri": "https://oauth2.googleapis.com/token",
            "redirect_uris": ["http://localhost:5000/auth/callback"]
        }
    }

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Azure Speech setup
speech_config = None
try:
    if AZURE_SPEECH_KEY and AZURE_SPEECH_REGION:
        speech_config = speechsdk.SpeechConfig(subscription=AZURE_SPEECH_KEY, region=AZURE_SPEECH_REGION)
        if AZURE_SPEECH_ENDPOINT:
            speech_config.endpoint_id = AZURE_SPEECH_ENDPOINT
        speech_config.speech_recognition_language = "en-US"
        speech_config.output_format = speechsdk.OutputFormat.Detailed
        print("Azure Speech client initialized successfully.")
    else:
        print("Warning: AZURE_SPEECH_KEY or AZURE_SPEECH_REGION not set.")
except Exception as e:
    print(f"Warning: Azure Speech client initialization failed: {e}")
    speech_config = None

# Authentication decorators
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            return jsonify({"error": "Authentication required", "redirect": "/auth/login"}), 401
        return f(*args, **kwargs)
    return decorated_function

def subscription_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            return jsonify({"error": "Authentication required", "redirect": "/auth/login"}), 401
        
        if not is_subscription_active(session['user_db']['id']):
            return jsonify({
                "error": "Active subscription required", 
                "redirect": "/subscription",
                "subscription_status": "inactive"
            }), 402  # Payment Required
        return f(*args, **kwargs)
    return decorated_function

# Helper functions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def convert_to_wav(input_path, output_path):
    try:
        audio = AudioSegment.from_file(input_path)
        audio = audio.set_channels(1).set_frame_rate(16000).set_sample_width(2)
        audio.export(output_path, format="wav")
    except Exception as e:
        print(f"Audio conversion error: {e}")
        return False
    return True

def transcribe_audio(audio_path):
    if not speech_config:
        return {"error": "Azure Speech client not initialized."}
    
    try:
        audio_config = speechsdk.audio.AudioConfig(filename=audio_path)
        speech_recognizer = speechsdk.SpeechRecognizer(speech_config=speech_config, audio_config=audio_config)
        
        done = False
        results = []
        
        def stop_cb(evt):
            nonlocal done
            done = True
        
        def recognized_cb(evt):
            if evt.result.reason == speechsdk.ResultReason.RecognizedSpeech:
                results.append({
                    'text': evt.result.text,
                    'confidence': evt.result.properties.get(speechsdk.PropertyId.SpeechServiceResponse_JsonResult, '{}')
                })
        
        speech_recognizer.recognized.connect(recognized_cb)
        speech_recognizer.session_stopped.connect(stop_cb)
        speech_recognizer.canceled.connect(stop_cb)
        
        speech_recognizer.start_continuous_recognition()
        
        import time
        timeout = 300
        start_time = time.time()
        
        while not done and (time.time() - start_time) < timeout:
            time.sleep(0.5)
        
        speech_recognizer.stop_continuous_recognition()
        
        if not results:
            return {"error": "No speech recognized in the audio file."}
        
        full_transcript = " ".join([result['text'] for result in results])
        word_count = len(full_transcript.split()) if full_transcript else 0
        
        return {
            "transcript": full_transcript,
            "confidence": 0.85,
            "word_count": word_count
        }
        
    except Exception as e:
        return {"error": f"Transcription failed: {str(e)}"}

def create_docx(transcript, filename="transcript"):
    doc = Document()
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = f"Audio Transcript - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    
    title = doc.add_heading('Audio Transcription', 0)
    doc.add_paragraph(transcript)
    
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Word count: {len(transcript.split())} words"
    
    temp_path = os.path.join(tempfile.gettempdir(), f"{filename}_{uuid.uuid4().hex}.docx")
    doc.save(temp_path)
    return temp_path

# Routes
@app.route('/')
def index():
    return send_file('index.html')

@app.route('/api/status')
def api_status():
    user_status = {}
    if 'user' in session and 'user_db' in session:
        user_status = {
            "authenticated": True,
            "user": session['user']['name'],
            "subscription_active": is_subscription_active(session['user_db']['id']),
            "subscription_status": session['user_db']['subscription_status']
        }
    else:
        user_status = {"authenticated": False}
    
    return jsonify({
        "message": "Speech-to-Text API is running",
        **user_status
    })

# Subscription status endpoint
@app.route('/api/subscription/status')
@login_required
def subscription_status():
    """Get current user's subscription status."""
    try:
        user_db = session['user_db']
        subscription_active = is_subscription_active(user_db['id'])
        
        return jsonify({
            "subscription_active": subscription_active,
            "subscription_status": user_db.get('subscription_status', 'inactive'),
            "user_id": user_db['id']
        })
    except Exception as e:
        print(f"Error checking subscription status: {e}")
        return jsonify({"error": "Failed to check subscription status"}), 500

# Subscription success/cancel pages
@app.route('/subscription/success')
def subscription_success():
    """Handle successful subscription."""
    return redirect('/?subscription=success')

@app.route('/subscription/cancel') 
def subscription_cancel():
    """Handle cancelled subscription."""
    return redirect('/?subscription=cancelled')

# Customer portal endpoint
@app.route('/api/subscription/portal', methods=['POST'])
@login_required
def create_portal_session():
    """Create Stripe customer portal session."""
    try:
        user_db = session['user_db']
        
        if not user_db.get('stripe_customer_id'):
            return jsonify({"error": "No Stripe customer found"}), 400
        
        portal_session = stripe.billing_portal.Session.create(
            customer=user_db['stripe_customer_id'],
            return_url=url_for('index', _external=True)
        )
        
        return jsonify({"portal_url": portal_session.url})
        
    except Exception as e:
        print(f"Portal session creation failed: {e}")
        return jsonify({"error": str(e)}), 500

# Stripe webhook endpoint
@app.route('/webhook/stripe', methods=['POST'])
def stripe_webhook():
    """Handle Stripe webhook events."""
    payload = request.get_data()
    sig_header = request.headers.get('Stripe-Signature')
    
    if not STRIPE_WEBHOOK_SECRET:
        print("Warning: STRIPE_WEBHOOK_SECRET not configured")
        return jsonify({"error": "Webhook secret not configured"}), 400
    
    try:
        event = stripe.Webhook.construct_event(
            payload, sig_header, STRIPE_WEBHOOK_SECRET
        )
    except ValueError as e:
        print(f"Invalid payload: {e}")
        return jsonify({"error": "Invalid payload"}), 400
    except stripe.error.SignatureVerificationError as e:
        print(f"Invalid signature: {e}")
        return jsonify({"error": "Invalid signature"}), 400
    
    # Handle the event
    try:
        if event['type'] == 'customer.subscription.created':
            handle_subscription_created(event['data']['object'])
        elif event['type'] == 'customer.subscription.updated':
            handle_subscription_updated(event['data']['object'])
        elif event['type'] == 'customer.subscription.deleted':
            handle_subscription_deleted(event['data']['object'])
        elif event['type'] == 'invoice.payment_succeeded':
            handle_payment_succeeded(event['data']['object'])
        elif event['type'] == 'invoice.payment_failed':
            handle_payment_failed(event['data']['object'])
        else:
            print(f"Unhandled event type: {event['type']}")
        
        return jsonify({"status": "success"}), 200
        
    except Exception as e:
        print(f"Error handling webhook event {event['type']}: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": "Webhook handler failed"}), 500

# Debug endpoints
@app.route('/debug/user-info')
@login_required
def debug_user_info():
    user_db = session['user_db']
    
    # Get detailed user info from database
    with get_db_connection() as conn:
        user = execute_query(conn,
            'SELECT * FROM users WHERE id = ?',
            (user_db['id'],),
            fetch='one'
        )
        
        if user:
            user_dict = dict(user)
            # Don't expose sensitive data in debug
            return jsonify({
                "user_id": user_dict['id'],
                "email": user_dict['email'],
                "stripe_customer_id": user_dict['stripe_customer_id'],
                "subscription_status": user_dict['subscription_status'],
                "subscription_id": user_dict['subscription_id'],
                "subscription_end_date": str(user_dict['subscription_end_date']) if user_dict['subscription_end_date'] else None,
                "created_at": str(user_dict['created_at']),
                "updated_at": str(user_dict['updated_at'])
            })
        else:
            return jsonify({"error": "User not found in database"})

@app.route('/debug/subscription-info')
@login_required
def debug_subscription_info():
    """Debug endpoint to check subscription information."""
    user_db = session['user_db']
    
    try:
        # Get user from database
        with get_db_connection() as conn:
            user = execute_query(conn,
                'SELECT * FROM users WHERE id = ?',
                (user_db['id'],),
                fetch='one'
            )
        
        if not user:
            return jsonify({"error": "User not found"}), 404
        
        user_dict = dict(user)
        subscription_active = is_subscription_active(user_db['id'])
        
        # Try to get Stripe customer info if available
        stripe_info = None
        if user_dict.get('stripe_customer_id'):
            try:
                stripe_customer = stripe.Customer.retrieve(user_dict['stripe_customer_id'])
                stripe_info = {
                    "id": stripe_customer.id,
                    "email": stripe_customer.email,
                    "created": stripe_customer.created
                }
                
                # Get subscriptions
                subscriptions = stripe.Subscription.list(customer=stripe_customer.id)
                stripe_info["subscriptions"] = [
                    {
                        "id": sub.id,
                        "status": sub.status,
                        "current_period_end": sub.current_period_end
                    } for sub in subscriptions.data
                ]
                
            except Exception as e:
                stripe_info = {"error": str(e)}
        
        return jsonify({
            "user_id": user_dict['id'],
            "email": user_dict['email'],
            "stripe_customer_id": user_dict.get('stripe_customer_id'),
            "subscription_status": user_dict.get('subscription_status'),
            "subscription_id": user_dict.get('subscription_id'),
            "subscription_end_date": str(user_dict.get('subscription_end_date')) if user_dict.get('subscription_end_date') else None,
            "subscription_active_calculated": subscription_active,
            "stripe_info": stripe_info
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# Authentication routes
@app.route('/auth/login')
def login():
    if not client_config:
        return jsonify({"error": "Google OAuth not configured"}), 500
    
    try:
        google_scopes = [
            'openid',
            'https://www.googleapis.com/auth/userinfo.email',
            'https://www.googleapis.com/auth/userinfo.profile'
        ]
        
        flow = Flow.from_client_config(client_config, scopes=google_scopes)
        flow.redirect_uri = url_for('auth_callback', _external=True)
        
        authorization_url, state = flow.authorization_url(
            access_type='offline',
            include_granted_scopes='true'
        )
        
        session['state'] = state
        session['oauth_scopes'] = google_scopes
        
        return jsonify({"auth_url": authorization_url})
    
    except Exception as e:
        return jsonify({"error": f"Login failed: {str(e)}"}), 500

@app.route('/auth/callback')
def auth_callback():
    if not client_config:
        return jsonify({"error": "Google OAuth not configured"}), 500
    
    try:
        if 'state' not in session or request.args.get('state') != session['state']:
            return redirect('/?auth=error')
        
        flow = Flow.from_client_config(
            client_config,
            scopes=[
                'openid',
                'https://www.googleapis.com/auth/userinfo.email',
                'https://www.googleapis.com/auth/userinfo.profile'
            ],
            state=session['state']
        )
        flow.redirect_uri = url_for('auth_callback', _external=True)
        
        flow.fetch_token(authorization_response=request.url)
        
        credentials = flow.credentials
        request_session = google.auth.transport.requests.Request()
        
        id_info = id_token.verify_oauth2_token(
            credentials.id_token,
            request_session,
            GOOGLE_CLIENT_ID
        )
        
        # Store Google user info in session
        session['user'] = {
            'id': id_info['sub'],
            'email': id_info['email'],
            'name': id_info['name'],
            'picture': id_info.get('picture', ''),
            'verified_email': id_info.get('email_verified', False)
        }
        
        # Get or create user in database
        user_db = get_or_create_user(session['user'])
        session['user_db'] = user_db
        
        session.pop('state', None)
        session.pop('oauth_scopes', None)
        
        return redirect('/?auth=success')
        
    except Exception as e:
        print(f"Authentication error: {e}")
        session.pop('state', None)
        session.pop('oauth_scopes', None)
        return redirect('/?auth=error')

@app.route('/auth/logout', methods=['POST'])
def logout():
    session.clear()
    return jsonify({"message": "Logged out successfully"})

@app.route('/auth/user')
def get_user():
    if 'user' in session and 'user_db' in session:
        return jsonify({
            "authenticated": True,
            "user": session['user'],
            "subscription_active": is_subscription_active(session['user_db']['id']),
            "subscription_status": session['user_db']['subscription_status']
        })
    else:
        return jsonify({"authenticated": False, "user": None})

# Subscription routes
@app.route('/api/subscription/create-checkout-session', methods=['POST'])
@login_required
def create_checkout_session():
    try:
        # Always get the freshest user_db from the session at the start of the function
        # This ensures we start with the most current data, especially after a login.
        user_db = session['user_db'] 
        current_stripe_customer_id = user_db['stripe_customer_id']
        
        print(f"DEBUG: create_checkout_session - Initial customer_id from session: '{current_stripe_customer_id}'")
        print(f"DEBUG: Type of current_stripe_customer_id: {type(current_stripe_customer_id)}")

        customer_to_use_id = None # This will hold the final, valid Stripe customer ID

        if current_stripe_customer_id:
            try:
                # Attempt to retrieve existing customer to verify it's valid
                stripe.Customer.retrieve(current_stripe_customer_id)
                customer_to_use_id = current_stripe_customer_id # It's valid, so use it
                print(f"DEBUG: Stored Stripe customer '{current_stripe_customer_id}' found and is valid.")
            except stripe.error.InvalidRequestError as e:
                # Customer ID is invalid or not found in Stripe, so we'll create a new one
                print(f"DEBUG: Warning: Stored Stripe customer ID '{current_stripe_customer_id}' for user {user_db['id']} is invalid or not found in Stripe. Error: {e}. Will create a new customer.")
                
                # Clear the invalid ID from the database and session immediately
                update_user_subscription(user_db['id'], None, user_db['subscription_status'], user_db.get('subscription_id'), user_db.get('subscription_end_date'))
                # After updating, explicitly re-read user_db from session to ensure local variable is fresh
                user_db = session['user_db']
                customer_to_use_id = None # Explicitly set for clarity in this scope
            except Exception as e:
                # Catch any other unexpected errors during retrieval, also create new customer
                print(f"DEBUG: Unexpected error retrieving Stripe customer '{current_stripe_customer_id}' for user {user_db['id']}: {e}. Will create a new customer.")
                update_user_subscription(user_db['id'], None, user_db['subscription_status'], user_db.get('subscription_id'), user_db.get('subscription_end_date'))
                # After updating, explicitly re-read user_db from session to ensure local variable is fresh
                user_db = session['user_db']
                customer_to_use_id = None
        else:
            print(f"DEBUG: No Stripe customer ID found in session initially for user {user_db['id']}. Will create a new customer.")
            customer_to_use_id = None # Explicitly ensure it's None if not present

        # If customer_to_use_id is still None (either no customer_id initially, or it was invalid/failed retrieval)
        if not customer_to_use_id:
            print(f"DEBUG: Creating new Stripe customer for user {user_db['id']} ({session['user']['email']}).")
            customer = stripe.Customer.create(
                email=session['user']['email'],
                name=session['user']['name']
            )
            customer_to_use_id = customer.id
            print(f"DEBUG: New Stripe customer created: '{customer_to_use_id}'")
            
            # Update user's DB and session with the new customer ID
            update_user_subscription(user_db['id'], customer_to_use_id, user_db['subscription_status'])
            # After updating, explicitly re-read user_db from session to ensure local variable is fresh
            user_db = session['user_db']
            
        print(f"DEBUG: Final Stripe customer ID for checkout session: '{customer_to_use_id}'")
        # Create checkout session
        checkout_session = stripe.checkout.Session.create(
            customer=customer_to_use_id, # Use the determined final_customer_id
            payment_method_types=['card'],
            line_items=[{
                'price': SUBSCRIPTION_PRICE_ID,
                'quantity': 1,
            }],
            mode='subscription',
            success_url=url_for('subscription_success', _external=True), # Keep _external=True as Stripe requires absolute URLs
            cancel_url=url_for('subscription_cancel', _external=True),   # Keep _external=True as Stripe requires absolute URLs
            metadata={'user_id': str(user_db['id'])}
        )
        
        return jsonify({'checkout_url': checkout_session.url})
        
    except Exception as e:
        print(f"Checkout session creation failed: {e}")
        import traceback
        traceback.print_exc() # Print full traceback for debugging
        return jsonify({'error': str(e)}), 500

# Protected transcription routes
@app.route('/upload', methods=['POST'])
@subscription_required
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400
    
    if not allowed_file(file.filename):
        return jsonify({"error": f"File type not supported. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"}), 400
    
    try:
        filename = secure_filename(file.filename)
        unique_filename = f"{uuid.uuid4().hex}_{filename}"
        file_path = os.path.join(UPLOAD_FOLDER, unique_filename)
        file.save(file_path)
        
        wav_path = file_path
        if not filename.lower().endswith('.wav'):
            wav_path = os.path.join(UPLOAD_FOLDER, f"{unique_filename}.wav")
            if not convert_to_wav(file_path, wav_path):
                os.remove(file_path)
                return jsonify({"error": "Failed to convert audio file to WAV."}), 500
        
        result = transcribe_audio(wav_path)
        
        # Clean up files
        try:
            os.remove(file_path)
            if wav_path != file_path:
                os.remove(wav_path)
        except Exception:
            pass
        
        if "error" in result:
            return jsonify(result), 500
        
        # Save transcription to database
        user_db = session['user_db']
        with get_db_connection() as conn:
            execute_query(conn, '''
                INSERT INTO transcriptions (user_id, filename, transcript, word_count, confidence)
                VALUES (?, ?, ?, ?, ?)
            ''', (user_db['id'], filename, result['transcript'], result['word_count'], result['confidence']))
        
        return jsonify({
            "success": True,
            "transcript": result["transcript"],
            "confidence": result["confidence"],
            "word_count": result["word_count"],
            "filename": filename,
            "user": session['user']['name']
        })
    
    except Exception as e:
        print(f"Upload failed: {e}")
        return jsonify({"error": f"Processing failed: {str(e)}"}), 500

@app.route('/download', methods=['POST'])
@subscription_required
def download_transcript():
    data = request.get_json()
    if not data or 'transcript' not in data:
        return jsonify({"error": "No transcript provided"}), 400
    
    try:
        transcript = data['transcript']
        filename = data.get('filename', 'transcript')
        
        if '.' in filename:
            filename = filename.rsplit('.', 1)[0]
        
        docx_path = create_docx(transcript, filename)
        
        return send_file(
            docx_path,
            as_attachment=True,
            download_name=f"{filename}_transcript.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    except Exception as e:
        return jsonify({"error": f"Document creation failed: {str(e)}"}), 500

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({
        "status": "healthy",
        "database_type": "PostgreSQL" if (DATABASE_URL.startswith('postgresql://') and POSTGRES_AVAILABLE) else "SQLite",
        "azure_speech_initialized": speech_config is not None,
        "google_oauth_configured": client_config is not None,
        "stripe_configured": stripe.api_key is not None,
        "authenticated_session_active": 'user' in session,
        "timestamp": datetime.now().isoformat()
    })

# Add these database management routes to your app.py
#################################################################################################################DEBUGSTARTSHERE###########################################
@app.route('/debug/database/users')
@login_required
def debug_list_all_users():
    """List all users in the database with their Stripe info."""
    try:
        with get_db_connection() as conn:
            users = execute_query(conn, '''
                SELECT id, google_id, email, name, stripe_customer_id, 
                       subscription_status, subscription_id, subscription_end_date,
                       created_at, updated_at
                FROM users 
                ORDER BY created_at DESC
            ''', fetch='all')
            
            user_list = []
            for user in users:
                user_dict = dict(user)
                user_dict['created_at'] = str(user_dict['created_at']) if user_dict['created_at'] else None
                user_dict['updated_at'] = str(user_dict['updated_at']) if user_dict['updated_at'] else None
                user_dict['subscription_end_date'] = str(user_dict['subscription_end_date']) if user_dict['subscription_end_date'] else None
                
                # Check if Stripe customer exists
                stripe_status = "NOT_SET"
                if user_dict['stripe_customer_id']:
                    try:
                        stripe_customer = stripe.Customer.retrieve(user_dict['stripe_customer_id'])
                        if getattr(stripe_customer, 'deleted', False):
                            stripe_status = "DELETED"
                        else:
                            stripe_status = "EXISTS"
                    except stripe.error.InvalidRequestError:
                        stripe_status = "INVALID"
                    except Exception as e:
                        stripe_status = f"ERROR: {str(e)}"
                
                user_dict['stripe_customer_status'] = stripe_status
                user_list.append(user_dict)
            
            return jsonify({
                "total_users": len(user_list),
                "users": user_list
            })
            
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/debug/database/user/<int:user_id>')
@login_required
def debug_get_user_details(user_id):
    """Get detailed information about a specific user."""
    try:
        with get_db_connection() as conn:
            # Get user info
            user = execute_query(conn, '''
                SELECT * FROM users WHERE id = ?
            ''', (user_id,), fetch='one')
            
            if not user:
                return jsonify({"error": "User not found"}), 404
            
            user_dict = dict(user)
            user_dict['created_at'] = str(user_dict['created_at']) if user_dict['created_at'] else None
            user_dict['updated_at'] = str(user_dict['updated_at']) if user_dict['updated_at'] else None
            user_dict['subscription_end_date'] = str(user_dict['subscription_end_date']) if user_dict['subscription_end_date'] else None
            
            # Get transcription history
            transcriptions = execute_query(conn, '''
                SELECT id, filename, word_count, confidence, created_at
                FROM transcriptions 
                WHERE user_id = ? 
                ORDER BY created_at DESC
                LIMIT 10
            ''', (user_id,), fetch='all')
            
            transcription_list = []
            for trans in transcriptions:
                trans_dict = dict(trans)
                trans_dict['created_at'] = str(trans_dict['created_at']) if trans_dict['created_at'] else None
                transcription_list.append(trans_dict)
            
            # Check Stripe customer status
            stripe_info = {"status": "NOT_SET"}
            if user_dict['stripe_customer_id']:
                try:
                    stripe_customer = stripe.Customer.retrieve(user_dict['stripe_customer_id'])
                    stripe_info = {
                        "status": "EXISTS",
                        "id": stripe_customer.id,
                        "email": stripe_customer.email,
                        "created": stripe_customer.created,
                        "deleted": getattr(stripe_customer, 'deleted', False)
                    }
                    
                    # Get subscriptions for this customer
                    subscriptions = stripe.Subscription.list(customer=stripe_customer.id, limit=5)
                    stripe_info["subscriptions"] = [
                        {
                            "id": sub.id,
                            "status": sub.status,
                            "current_period_start": sub.current_period_start,
                            "current_period_end": sub.current_period_end,
                            "created": sub.created
                        } for sub in subscriptions.data
                    ]
                    
                except stripe.error.InvalidRequestError:
                    stripe_info = {"status": "INVALID", "error": "Customer not found in Stripe"}
                except Exception as e:
                    stripe_info = {"status": "ERROR", "error": str(e)}
            
            return jsonify({
                "user": user_dict,
                "transcriptions": {
                    "count": len(transcription_list),
                    "recent": transcription_list
                },
                "stripe_info": stripe_info,
                "subscription_active": is_subscription_active(user_id)
            })
            
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/debug/database/cleanup-invalid-customers', methods=['POST'])
@login_required
def debug_cleanup_invalid_customers():
    """Clean up users with invalid Stripe customer IDs."""
    try:
        with get_db_connection() as conn:
            # Get all users with Stripe customer IDs
            users = execute_query(conn, '''
                SELECT id, email, stripe_customer_id FROM users 
                WHERE stripe_customer_id IS NOT NULL AND stripe_customer_id != ''
            ''', fetch='all')
            
            cleaned_users = []
            errors = []
            
            for user in users:
                user_dict = dict(user)
                customer_id = user_dict['stripe_customer_id']
                
                try:
                    # Check if customer exists in Stripe
                    stripe_customer = stripe.Customer.retrieve(customer_id)
                    if getattr(stripe_customer, 'deleted', False):
                        # Customer is deleted, clean it up
                        execute_query(conn, '''
                            UPDATE users 
                            SET stripe_customer_id = NULL, 
                                subscription_status = 'inactive',
                                subscription_id = NULL,
                                subscription_end_date = NULL,
                                updated_at = CURRENT_TIMESTAMP
                            WHERE id = ?
                        ''', (user_dict['id'],))
                        
                        cleaned_users.append({
                            "user_id": user_dict['id'],
                            "email": user_dict['email'],
                            "customer_id": customer_id,
                            "reason": "Customer deleted in Stripe"
                        })
                        
                except stripe.error.InvalidRequestError:
                    # Customer doesn't exist, clean it up
                    execute_query(conn, '''
                        UPDATE users 
                        SET stripe_customer_id = NULL, 
                            subscription_status = 'inactive',
                            subscription_id = NULL,
                            subscription_end_date = NULL,
                            updated_at = CURRENT_TIMESTAMP
                        WHERE id = ?
                    ''', (user_dict['id'],))
                    
                    cleaned_users.append({
                        "user_id": user_dict['id'],
                        "email": user_dict['email'],
                        "customer_id": customer_id,
                        "reason": "Customer not found in Stripe"
                    })
                    
                except Exception as e:
                    errors.append({
                        "user_id": user_dict['id'],
                        "email": user_dict['email'],
                        "customer_id": customer_id,
                        "error": str(e)
                    })
            
            # Update current user's session if they were cleaned
            if 'user_db' in session:
                current_user_id = session['user_db']['id']
                if any(u['user_id'] == current_user_id for u in cleaned_users):
                    # Refresh session data
                    updated_user = execute_query(conn, 'SELECT * FROM users WHERE id = ?', (current_user_id,), fetch='one')
                    if updated_user:
                        session['user_db'] = dict(updated_user)
            
            return jsonify({
                "message": f"Cleanup completed. {len(cleaned_users)} users cleaned up.",
                "cleaned_users": cleaned_users,
                "errors": errors,
                "total_processed": len(users)
            })
            
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/debug/database/reset-user-stripe/<int:user_id>', methods=['POST'])
@login_required
def debug_reset_user_stripe(user_id):
    """Reset a specific user's Stripe information."""
    try:
        with get_db_connection() as conn:
            # Get user info
            user = execute_query(conn, '''
                SELECT id, email, name, stripe_customer_id FROM users WHERE id = ?
            ''', (user_id,), fetch='one')
            
            if not user:
                return jsonify({"error": "User not found"}), 404
            
            user_dict = dict(user)
            old_customer_id = user_dict['stripe_customer_id']
            
            # Clear Stripe data
            execute_query(conn, '''
                UPDATE users 
                SET stripe_customer_id = NULL, 
                    subscription_status = 'inactive',
                    subscription_id = NULL,
                    subscription_end_date = NULL,
                    updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
            ''', (user_id,))
            
            # Update session if this is the current user
            if 'user_db' in session and session['user_db']['id'] == user_id:
                updated_user = execute_query(conn, 'SELECT * FROM users WHERE id = ?', (user_id,), fetch='one')
                if updated_user:
                    session['user_db'] = dict(updated_user)
            
            return jsonify({
                "message": f"User {user_id} Stripe data reset successfully",
                "user_id": user_id,
                "email": user_dict['email'],
                "old_customer_id": old_customer_id,
                "new_status": "inactive"
            })
            
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/debug/database/stats')
@login_required
def debug_database_stats():
    """Get database statistics."""
    try:
        with get_db_connection() as conn:
            # User statistics
            total_users = execute_query(conn, 'SELECT COUNT(*) as count FROM users', fetch='one')['count']
            
            users_with_stripe = execute_query(conn, '''
                SELECT COUNT(*) as count FROM users 
                WHERE stripe_customer_id IS NOT NULL AND stripe_customer_id != ''
            ''', fetch='one')['count']
            
            active_subscriptions = execute_query(conn, '''
                SELECT COUNT(*) as count FROM users 
                WHERE subscription_status = 'active'
            ''', fetch='one')['count']
            
            # Transcription statistics
            total_transcriptions = execute_query(conn, 'SELECT COUNT(*) as count FROM transcriptions', fetch='one')['count']
            
            # Users by subscription status
            status_counts = execute_query(conn, '''
                SELECT subscription_status, COUNT(*) as count 
                FROM users 
                GROUP BY subscription_status
            ''', fetch='all')
            
            status_breakdown = {}
            for row in status_counts:
                status_breakdown[row['subscription_status'] or 'null'] = row['count']
            
            # Recent activity
            recent_users = execute_query(conn, '''
                SELECT COUNT(*) as count FROM users 
                WHERE created_at > CURRENT_TIMESTAMP - INTERVAL '7 days'
            ''', fetch='one') if DATABASE_URL.startswith('postgresql://') else execute_query(conn, '''
                SELECT COUNT(*) as count FROM users 
                WHERE created_at > datetime('now', '-7 days')
            ''', fetch='one')
            
            return jsonify({
                "users": {
                    "total": total_users,
                    "with_stripe_customer": users_with_stripe,
                    "active_subscriptions": active_subscriptions,
                    "recent_7_days": recent_users['count']
                },
                "transcriptions": {
                    "total": total_transcriptions
                },
                "subscription_status_breakdown": status_breakdown,
                "database_type": "PostgreSQL" if (DATABASE_URL.startswith('postgresql://') and POSTGRES_AVAILABLE) else "SQLite"
            })
            
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/debug/database/delete-user/<int:user_id>', methods=['DELETE'])
@login_required
def debug_delete_user(user_id):
    """Delete a user and all their data (use with caution!)."""
    try:
        # Don't allow users to delete themselves
        if 'user_db' in session and session['user_db']['id'] == user_id:
            return jsonify({"error": "Cannot delete your own account via debug endpoint"}), 400
        
        with get_db_connection() as conn:
            # Get user info first
            user = execute_query(conn, 'SELECT * FROM users WHERE id = ?', (user_id,), fetch='one')
            if not user:
                return jsonify({"error": "User not found"}), 404
            
            user_dict = dict(user)
            
            # Delete transcriptions first (foreign key constraint)
            transcription_count = execute_query(conn, 'SELECT COUNT(*) as count FROM transcriptions WHERE user_id = ?', (user_id,), fetch='one')['count']
            execute_query(conn, 'DELETE FROM transcriptions WHERE user_id = ?', (user_id,))
            
            # Delete user
            execute_query(conn, 'DELETE FROM users WHERE id = ?', (user_id,))
            
            return jsonify({
                "message": f"User {user_id} deleted successfully",
                "deleted_user": {
                    "id": user_dict['id'],
                    "email": user_dict['email'],
                    "name": user_dict['name']
                },
                "deleted_transcriptions": transcription_count
            })
            
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# Add a simple HTML interface for easier database management
@app.route('/debug/database/admin')
@login_required
def debug_database_admin():
    """Simple HTML interface for database management."""
    html = '''
    <!DOCTYPE html>
    <html>
    <head>
        <title>Database Admin</title>
        <style>
            body { font-family: Arial, sans-serif; margin: 20px; }
            .section { margin: 20px 0; padding: 15px; border: 1px solid #ddd; border-radius: 5px; }
            .button { background: #007bff; color: white; padding: 10px 15px; border: none; border-radius: 3px; cursor: pointer; margin: 5px; }
            .button:hover { background: #0056b3; }
            .danger { background: #dc3545; }
            .danger:hover { background: #c82333; }
            .success { background: #28a745; }
            .success:hover { background: #218838; }
            .output { background: #f8f9fa; padding: 10px; border-radius: 3px; margin: 10px 0; white-space: pre-wrap; }
        </style>
    </head>
    <body>
        <h1>Database Administration</h1>
        
        <div class="section">
            <h3>Database Overview</h3>
            <button class="button" onclick="loadStats()">Load Statistics</button>
            <button class="button" onclick="loadUsers()">List All Users</button>
            <div id="stats-output" class="output" style="display:none;"></div>
        </div>
        
        <div class="section">
            <h3>Stripe Cleanup</h3>
            <button class="button success" onclick="cleanupInvalidCustomers()">Cleanup Invalid Stripe Customers</button>
            <div id="cleanup-output" class="output" style="display:none;"></div>
        </div>
        
        <div class="section">
            <h3>User Management</h3>
            <input type="number" id="user-id" placeholder="User ID" />
            <button class="button" onclick="getUserDetails()">Get User Details</button>
            <button class="button" onclick="resetUserStripe()">Reset User Stripe</button>
            <button class="button danger" onclick="deleteUser()">Delete User</button>
            <div id="user-output" class="output" style="display:none;"></div>
        </div>
        
        <script>
            async function loadStats() {
                try {
                    const response = await fetch('/debug/database/stats');
                    const data = await response.json();
                    document.getElementById('stats-output').style.display = 'block';
                    document.getElementById('stats-output').textContent = JSON.stringify(data, null, 2);
                } catch (error) {
                    alert('Error: ' + error.message);
                }
            }
            
            async function loadUsers() {
                try {
                    const response = await fetch('/debug/database/users');
                    const data = await response.json();
                    document.getElementById('stats-output').style.display = 'block';
                    document.getElementById('stats-output').textContent = JSON.stringify(data, null, 2);
                } catch (error) {
                    alert('Error: ' + error.message);
                }
            }
            
            async function cleanupInvalidCustomers() {
                if (!confirm('This will clean up all invalid Stripe customer references. Continue?')) return;
                try {
                    const response = await fetch('/debug/database/cleanup-invalid-customers', {method: 'POST'});
                    const data = await response.json();
                    document.getElementById('cleanup-output').style.display = 'block';
                    document.getElementById('cleanup-output').textContent = JSON.stringify(data, null, 2);
                } catch (error) {
                    alert('Error: ' + error.message);
                }
            }
            
            async function getUserDetails() {
                const userId = document.getElementById('user-id').value;
                if (!userId) { alert('Please enter a user ID'); return; }
                try {
                    const response = await fetch(`/debug/database/user/${userId}`);
                    const data = await response.json();
                    document.getElementById('user-output').style.display = 'block';
                    document.getElementById('user-output').textContent = JSON.stringify(data, null, 2);
                } catch (error) {
                    alert('Error: ' + error.message);
                }
            }
            
            async function resetUserStripe() {
                const userId = document.getElementById('user-id').value;
                if (!userId) { alert('Please enter a user ID'); return; }
                if (!confirm(`Reset Stripe data for user ${userId}?`)) return;
                try {
                    const response = await fetch(`/debug/database/reset-user-stripe/${userId}`, {method: 'POST'});
                    const data = await response.json();
                    document.getElementById('user-output').style.display = 'block';
                    document.getElementById('user-output').textContent = JSON.stringify(data, null, 2);
                } catch (error) {
                    alert('Error: ' + error.message);
                }
            }
            
            async function deleteUser() {
                const userId = document.getElementById('user-id').value;
                if (!userId) { alert('Please enter a user ID'); return; }
                if (!confirm(`DELETE user ${userId}? This cannot be undone!`)) return;
                if (!confirm(`Are you SURE you want to delete user ${userId}?`)) return;
                try {
                    const response = await fetch(`/debug/database/delete-user/${userId}`, {method: 'DELETE'});
                    const data = await response.json();
                    document.getElementById('user-output').style.display = 'block';
                    document.getElementById('user-output').textContent = JSON.stringify(data, null, 2);
                } catch (error) {
                    alert('Error: ' + error.message);
                }
            }
        </script>
    </body>
    </html>
    '''
    return html
# Initialize database on startup
init_db()

# Print the URL map to debug registered routes
print("\n--- Flask URL Map ---")
for rule in app.url_map.iter_rules():
    print(f"Endpoint: {rule.endpoint}, Methods: {rule.methods}, Rule: {rule.rule}")
print("---------------------\n")

if __name__ == '__main__':
    # This block is for local development only
    print("Running Flask app in local development mode...")
    app.run(debug=True, host='0.0.0.0', port=5000)
else:
    # This block is for production deployment (e.g., with Gunicorn)
    # The 'app' object is expected to be imported by a WSGI server
    print("Flask app loaded for production deployment.")
